from flask import Flask, request, send_file, abort
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

pdfmetrics.registerFont(TTFont('STHeiti', '/System/Library/Fonts/STHeiti Medium.ttc', subfontIndex=0))

app = Flask(__name__, static_folder='.', static_url_path='')

ZONES = [
    {'id': 'z1', 'name': '空間一', 'subtitle': '一樓大廳',     'desc': '開幕式・藝文表演'},
    {'id': 'z2', 'name': '空間二', 'subtitle': '二樓特展室',   'desc': '靜態展（海報・數位藝術・展品）'},
    {'id': 'z3', 'name': '空間三', 'subtitle': '二樓電梯口',   'desc': '動態展A區'},
    {'id': 'z4', 'name': '空間四', 'subtitle': '二樓手扶梯口', 'desc': '動態展B區（在地文化）'},
    {'id': 'z5', 'name': '空間五', 'subtitle': '三樓手扶梯口', 'desc': '動態展C區（VR等）'},
    {'id': 'z6', 'name': '空間六', 'subtitle': '會議廳',       'desc': '成果展典禮'},
    {'id': 'z7', 'name': '空間七', 'subtitle': '走道',         'desc': '穿梭空間'},
]


def hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def build_summary_html(exhibits, schools_map):
    now = datetime.now().strftime('%Y/%m/%d %H:%M')
    rows = ''
    for zone in ZONES:
        items = [e for e in exhibits if e.get('zone') == zone['id']]
        rows += f'''<tr>
            <td colspan="4" class="zone-header">
                {zone["name"]}：{zone["subtitle"]}　{zone["desc"]}
            </td>
        </tr>'''
        if not items:
            rows += '<tr><td colspan="4" class="empty">尚無展品分配</td></tr>'
        else:
            for ex in items:
                s = schools_map.get(ex.get('school', ''), {})
                color = s.get('color', '#666')
                name = s.get('name', ex.get('school', ''))
                status = '待確認' if ex.get('status') == 'pending' else '已確認'
                status_color = '#e65100' if ex.get('status') == 'pending' else '#27ae60'
                rows += f'''<tr>
                    <td style="color:{color};font-weight:700;white-space:nowrap">{name}</td>
                    <td>{ex.get('name', '')}</td>
                    <td>{ex.get('equipment', '') or '—'}</td>
                    <td style="color:{status_color}">{status}</td>
                </tr>'''

    pool_items = [e for e in exhibits if e.get('zone') == 'pool']
    if pool_items:
        rows += f'<tr><td colspan="4" class="pool-header">未分配展品（{len(pool_items)} 件）</td></tr>'
        for ex in pool_items:
            s = schools_map.get(ex.get('school', ''), {})
            color = s.get('color', '#666')
            name = s.get('name', ex.get('school', ''))
            status = '待確認' if ex.get('status') == 'pending' else '已確認'
            status_color = '#e65100' if ex.get('status') == 'pending' else '#27ae60'
            rows += f'''<tr>
                <td style="color:{color};font-weight:700;white-space:nowrap">{name}</td>
                <td>{ex.get('name', '')}</td>
                <td>{ex.get('equipment', '') or '—'}</td>
                <td style="color:{status_color}">{status}</td>
            </tr>'''

    return f'''<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{
    font-family: 'PingFang TC', 'Heiti TC', 'Microsoft JhengHei', Arial, sans-serif;
    padding: 32px; color: #1C1008; font-size: 13px;
  }}
  h1 {{ font-size: 18px; color: #4A2E18; margin: 0 0 4px; }}
  .subtitle {{ font-size: 11px; color: #888; margin: 0 0 20px; }}
  table {{ width: 100%; border-collapse: collapse; }}
  td, th {{ border: 1px solid #E5DDD3; padding: 7px 10px; }}
  th {{ background: #4A2E18; color: white; font-weight: 700; text-align: left; }}
  .zone-header {{ background: #F2EAE0; font-weight: 700; color: #4A2E18; }}
  .pool-header {{ background: #FEF3E8; font-weight: 700; color: #c0392b; }}
  .empty {{ color: #bbb; font-style: italic; }}
  tr:nth-child(even) td {{ background: #FAFAF8; }}
  tr:nth-child(odd) td {{ background: #FFFFFF; }}
  .zone-header, .pool-header, .empty {{ background: unset; }}
  tr:has(.zone-header) td, tr:has(.pool-header) td, tr:has(.empty) td {{ background: unset !important; }}
</style>
</head><body>
<h1>王船館展區規劃總覽</h1>
<p class="subtitle">青陽計畫 · 數位人文薈萃特展　匯出時間：{now}</p>
<table>
  <thead><tr><th>學校</th><th>展品名稱</th><th>備注</th><th>狀態</th></tr></thead>
  <tbody>{rows}</tbody>
</table>
</body></html>'''


@app.route('/')
def index():
    return app.send_static_file('index.html')


@app.route('/export/pdf', methods=['POST'])
def export_pdf():
    data = request.get_json()
    if not data:
        abort(400)
    exhibits = data.get('exhibits', [])
    schools_map = {s['id']: s for s in data.get('schools', [])}

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    font = 'STHeiti'
    title_style = ParagraphStyle('title', fontName=font, fontSize=16, spaceAfter=4, textColor=colors.HexColor('#4A2E18'))
    sub_style   = ParagraphStyle('sub',   fontName=font, fontSize=9,  spaceAfter=16, textColor=colors.HexColor('#888888'))
    body_style  = ParagraphStyle('body',  fontName=font, fontSize=9)

    story = [
        Paragraph('王船館展區規劃總覽', title_style),
        Paragraph(f'青陽計畫・數位人文薈萃特展　匯出時間：{datetime.now().strftime("%Y/%m/%d %H:%M")}', sub_style),
    ]

    col_widths = [2.5*cm, 9*cm, 4.5*cm, 1.8*cm]
    header_row = [Paragraph(h, ParagraphStyle('h', fontName=font, fontSize=9, textColor=colors.white))
                  for h in ['學校', '展品名稱', '備注', '狀態']]

    def make_table(items):
        rows = [header_row]
        for ex in items:
            s = schools_map.get(ex.get('school', ''), {})
            school_color = colors.HexColor(s.get('color', '#666666'))
            status = '待確認' if ex.get('status') == 'pending' else '已確認'
            status_color = colors.HexColor('#e65100') if ex.get('status') == 'pending' else colors.HexColor('#27ae60')
            rows.append([
                Paragraph(s.get('name', ''), ParagraphStyle('s', fontName=font, fontSize=9, textColor=school_color)),
                Paragraph(ex.get('name', ''), body_style),
                Paragraph(ex.get('equipment', '') or '—', body_style),
                Paragraph(status, ParagraphStyle('st', fontName=font, fontSize=9, textColor=status_color)),
            ])
        t = Table(rows, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4A2E18')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#FAFAF8')]),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E5DDD3')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        return t

    for zone in ZONES:
        items = [e for e in exhibits if e.get('zone') == zone['id']]
        zone_label = f'{zone["name"]}：{zone["subtitle"]}　{zone["desc"]}'
        zone_row = [[Paragraph(zone_label, ParagraphStyle('z', fontName=font, fontSize=9,
                               fontWeight='bold', textColor=colors.HexColor('#4A2E18')))]]
        zone_table = Table(zone_row, colWidths=[sum(col_widths)])
        zone_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F2EAE0')),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(zone_table)
        if items:
            story.append(make_table(items))
        else:
            story.append(Paragraph('　尚無展品分配', ParagraphStyle('e', fontName=font, fontSize=9,
                                   textColor=colors.HexColor('#bbbbbb'))))
        story.append(Spacer(1, 6))

    pool_items = [e for e in exhibits if e.get('zone') == 'pool']
    if pool_items:
        pool_label = f'未分配展品（{len(pool_items)} 件）'
        pool_row = [[Paragraph(pool_label, ParagraphStyle('p', fontName=font, fontSize=9,
                               textColor=colors.HexColor('#c0392b')))]]
        pool_table = Table(pool_row, colWidths=[sum(col_widths)])
        pool_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FEF3E8')),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(pool_table)
        story.append(make_table(pool_items))

    doc.build(story)
    buf.seek(0)
    return send_file(buf, mimetype='application/pdf',
                     as_attachment=True, download_name='王船館展區規劃總覽.pdf')


@app.route('/export/word', methods=['POST'])
def export_word():
    data = request.get_json()
    if not data:
        abort(400)
    exhibits = data.get('exhibits', [])
    schools_map = {s['id']: s for s in data.get('schools', [])}

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    title = doc.add_heading('王船館展區規劃總覽', 0)
    title.alignment = 1
    sub = doc.add_paragraph(
        f'青陽計畫 · 數位人文薈萃特展　匯出時間：{datetime.now().strftime("%Y/%m/%d %H:%M")}'
    )
    sub.alignment = 1
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x87, 0x86)
    doc.add_paragraph()

    def add_table(items):
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, text in enumerate(['學校', '展品名稱', '備注', '狀態']):
            hdr[i].text = text
            hdr[i].paragraphs[0].runs[0].bold = True
            tc = hdr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '4A2E18')
            tcPr.append(shd)
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for ex in items:
            s = schools_map.get(ex.get('school', ''), {})
            name = s.get('name', ex.get('school', ''))
            status = '待確認' if ex.get('status') == 'pending' else '已確認'
            row = table.add_row().cells
            row[0].text = name
            row[1].text = ex.get('name', '')
            row[2].text = ex.get('equipment', '') or '—'
            row[3].text = status
            color_hex = s.get('color', '#666666').lstrip('#')
            if len(color_hex) == 6:
                r, g, b = hex_to_rgb('#' + color_hex)
                row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(r, g, b)
        return table

    for zone in ZONES:
        items = [e for e in exhibits if e.get('zone') == zone['id']]
        heading = doc.add_heading(f'{zone["name"]}：{zone["subtitle"]}', 2)
        doc.add_paragraph(zone['desc']).runs[0].font.color.rgb = RGBColor(0x88, 0x87, 0x86)
        if not items:
            doc.add_paragraph('尚無展品分配').runs[0].italic = True
        else:
            add_table(items)
        doc.add_paragraph()

    pool_items = [e for e in exhibits if e.get('zone') == 'pool']
    if pool_items:
        doc.add_heading(f'未分配展品（{len(pool_items)} 件）', 2)
        add_table(pool_items)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='王船館展區規劃總覽.docx'
    )


if __name__ == '__main__':
    app.run(port=3456, debug=False)
